import docx, os

# Current Working Directory
print(os.getcwd())

# create document object with docx
doc = docx.Document('demo.docx')

# return paragraph objects in document object
print(doc.paragraphs)

# return text from first paragraph object
print(doc.paragraphs[0].text)

# save text seconds paragraph in variable
para = doc.paragraphs[1]

# RUNS == change in styles like bold or italic
print(para.runs)

# print first run in second paragraph
print(para.runs[0].text)

print(para.runs[1].text)
print(para.runs[2].text)
print(para.runs[3].text)

# check for bold style in 2nd run
print(para.runs[1].bold)

# check for italic style in 4th run
print(para.runs[3].italic)

# we can change styles in runs
para.runs[0].underline == True

# change text
para.runs[0].text = 'A underlined paragraph having some'
doc.save('demo.docx')       # better practice to save to new tmp document name => bugs

print(para.runs[0].text)

# paragraph and run object have a style attribute
para.style = 'Title'
doc.save('demo.docx')

# create new word document object
newDoc = docx.Document()

# add two paragraphs
newDoc.add_paragraph('Hello this is my first paragraph. ')
newDoc.add_paragraph('Hello this is my second paragraph. ')

# save first paragraph to newPara variable
newPara = newDoc.paragraphs[0]

# add run to first paragraph
newPara.add_run('This is a new run')

# set new run to bold
newPara.runs[1].bold = True

newDoc.save('newPythonDocx.docx')

# function to return entire text in python
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('demo.docx'))